from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_DIR = Path("researchdraft/outputs")
REQUIRED_FILES = [
    "draft_context.json",
    "draft.md",
    "paper.docx",
    "quality_report.md",
    "trace.json",
    "candidate_literature.json",
    "source_review_report.json",
    "search_cache.json",
]
REQUIRED_HEADINGS = [
    "# ",
    "## 摘要",
    "## 关键词",
    "## 引言",
    "## 相关工作",
    "## 方法",
    "## 实验与结果分析",
    "## 结论",
    "## 参考文献",
]
TRACE_KEYS = {
    "task_id",
    "agent",
    "stage",
    "input_keys",
    "output_keys",
    "tool_call",
    "status",
    "failure_reason",
    "timestamp",
}


def main() -> int:
    errors: list[str] = []
    for name in REQUIRED_FILES:
        path = OUTPUT_DIR / name
        if not path.exists() or path.stat().st_size == 0:
            errors.append(f"缺失或为空: {path}")

    draft_path = OUTPUT_DIR / "draft.md"
    if draft_path.exists():
        draft = draft_path.read_text("utf-8")
        for heading in REQUIRED_HEADINGS:
            if heading not in draft:
                errors.append(f"draft.md 缺少章节: {heading}")
        if "[待补充：" not in draft and "[待确认：" not in draft:
            errors.append("draft.md 未包含待补充/待确认标记")

    report_path = OUTPUT_DIR / "quality_report.md"
    if report_path.exists():
        report = report_path.read_text("utf-8")
        for heading in [
            "## 项目摘要",
            "## Draft Context 摘要",
            "## 结构检查",
            "## 内容缺失",
            "## 待补充项",
            "## 待确认项",
            "## 引用与文献质量检查",
            "## 联网文献搜索与人工确认",
            "## Word 输出检查",
            "## Agent 执行摘要",
            "## 版本限制",
        ]:
            if heading not in report:
                errors.append(f"quality_report.md 缺少模块: {heading}")
        for label in [
            "搜索 provider",
            "搜索 query 列表",
            "是否使用缓存",
            "候选去重前数量",
            "候选去重后数量",
            "高置信候选数量",
            "低置信候选数量",
            "来源类型分布",
            "搜索失败与 fallback 情况",
        ]:
            if label not in report:
                errors.append(f"quality_report.md 缺少搜索报告字段: {label}")

    trace_path = OUTPUT_DIR / "trace.json"
    if trace_path.exists():
        trace_items = json.loads(trace_path.read_text("utf-8"))
        for index, item in enumerate(trace_items, 1):
            if set(item) != TRACE_KEYS:
                errors.append(f"trace.json 第 {index} 步字段不完整")
        agents = {item.get("agent") for item in trace_items}
        for agent in (
            "InterviewAgent",
            "PlanningAgent",
            "WritingAgent",
            "LiteratureSearchAgent",
            "SourceReviewAgent",
            "HumanReviewGate",
            "LiteratureAgent",
            "CitationAgent",
            "WordFormatAgent",
            "VerifierAgent",
        ):
            if agent not in agents:
                errors.append(f"trace.json 缺少 {agent}")
        literature_entries = [item for item in trace_items if item.get("agent") == "LiteratureSearchAgent"]
        if not literature_entries:
            errors.append("trace.json 缺少 LiteratureSearchAgent 明细")
        else:
            tool_call = literature_entries[0].get("tool_call", "")
            for label in (
                "provider=",
                "queries=",
                "raw_result_count=",
                "deduped_result_count=",
                "cache_hit=",
                "fallback_used=",
            ):
                if label not in tool_call:
                    errors.append(f"LiteratureSearchAgent trace 缺少 {label}")

    candidate_path = OUTPUT_DIR / "candidate_literature.json"
    if candidate_path.exists():
        payload = json.loads(candidate_path.read_text("utf-8"))
        for key in (
            "provider",
            "queries",
            "raw_result_count",
            "deduped_result_count",
            "cache_hit",
            "fallback_used",
            "candidates",
        ):
            if key not in payload:
                errors.append(f"candidate_literature.json 缺少字段: {key}")
        for candidate in payload.get("candidates", []):
            for key in (
                "candidate_id",
                "title",
                "source_url",
                "snippet",
                "year",
                "source_type",
                "possible_venue",
                "is_academic_like",
                "risk_flags",
                "provider",
                "query",
                "confidence",
                "status",
            ):
                if key not in candidate:
                    errors.append(f"candidate_literature.json 候选缺少字段: {key}")
            if candidate.get("status") != "pending_review":
                errors.append("candidate_literature.json 候选状态不是 pending_review")

    docx_path = OUTPUT_DIR / "paper.docx"
    if docx_path.exists():
        doc = Document(docx_path)
        if not doc.paragraphs:
            errors.append("paper.docx 无段落")
        else:
            title = doc.paragraphs[0]
            if title.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                errors.append("paper.docx 标题未居中")
            if not any(run.bold for run in title.runs):
                errors.append("paper.docx 标题未加粗")
        if not any(p.text.strip() == "参考文献" for p in doc.paragraphs):
            errors.append("paper.docx 缺少参考文献节")

    if errors:
        for error in errors:
            print(f"FAIL: {error}")
        return 1

    print("ResearchDraft smoke check passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
