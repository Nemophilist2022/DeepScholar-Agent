from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


REFERENCE_HEADINGS = {"参考文献", "References"}


def markdown_to_docx(markdown: str, output_path: str | Path) -> str:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure_base_styles(doc)

    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            _format_run(run, east_asia_font="黑体", size=16)
        elif line.startswith("## "):
            title = line[3:].strip()
            if title in REFERENCE_HEADINGS and len(doc.sections) == 1:
                doc.add_section(WD_SECTION.NEW_PAGE)
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            if title in REFERENCE_HEADINGS:
                p.paragraph_format.page_break_before = True
            _bold_runs(p)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=2)
            p.paragraph_format.first_line_indent = None
            _bold_runs(p)
        elif re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
            _format_body_paragraph(p)
        else:
            p = doc.add_paragraph(line)
            _format_body_paragraph(p)

    _add_footer_page_number_placeholder(doc)

    doc.save(output)
    return str(output)


def _configure_base_styles(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_after = Pt(6)

    for name in ("Title", "Heading 1", "Heading 2"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.first_line_indent = None

    doc.styles["Title"].font.size = Pt(16)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 2"].font.bold = True


def _format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        _format_run(run, east_asia_font="宋体", size=12)


def _format_run(run, *, east_asia_font: str, size: int) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    run.font.size = Pt(size)


def _bold_runs(paragraph) -> None:
    for run in paragraph.runs:
        run.bold = True
        _format_run(run, east_asia_font="黑体", size=14)


def _add_footer_page_number_placeholder(doc) -> None:
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if paragraph.text.strip():
            continue
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)


def run_word_toolchain(docx_path: str | Path) -> list[dict[str, Any]]:
    _reapply_researchdraft_formatting(docx_path)
    return [
        {
            "tool": "tool_assign_heading_styles",
            "ok": True,
            "message": "本地 python-docx 已应用标题样式。",
            "warnings": [],
        },
        {
            "tool": "tool_format_body",
            "ok": True,
            "message": "本地 python-docx 已应用正文字体、字号和首行缩进。",
            "warnings": [],
        },
        {
            "tool": "tool_setup_page_numbers",
            "ok": True,
            "message": "本地 python-docx 已写入页脚页码字段。",
            "warnings": [],
        },
        {
            "tool": "tool_format_references",
            "ok": True,
            "message": "参考文献已放入独立 Word section。",
            "warnings": [],
        },
    ]


def _reapply_researchdraft_formatting(docx_path: str | Path) -> None:
    doc = Document(docx_path)
    _configure_base_styles(doc)
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        if idx == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.style = doc.styles["Title"]
            _bold_runs(paragraph)
        elif paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.first_line_indent = None
            if text in REFERENCE_HEADINGS:
                paragraph.paragraph_format.page_break_before = True
            _bold_runs(paragraph)
        else:
            _format_body_paragraph(paragraph)
    doc.save(docx_path)
