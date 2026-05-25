from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


REQUIRED_SECTIONS = [
    "摘要",
    "关键词",
    "引言",
    "相关工作",
    "方法",
    "实验与结果分析",
    "结论",
    "参考文献",
]

TRACE_REQUIRED_KEYS = {
    "task_id",
    "agent",
    "stage",
    "input_keys",
    "output_keys",
    "tool_call",
    "status",
    "failure_reason",
    "timestamp",
}


@dataclass
class FormatCheck:
    name: str
    passed: bool
    evidence: str


@dataclass
class VerificationResult:
    completed: list[str] = field(default_factory=list)
    missing_items: list[str] = field(default_factory=list)
    confirmation_items: list[str] = field(default_factory=list)
    structure_checks: list[FormatCheck] = field(default_factory=list)
    file_checks: list[FormatCheck] = field(default_factory=list)
    trace_checks: list[FormatCheck] = field(default_factory=list)
    format_checks: list[FormatCheck] = field(default_factory=list)
    citation_report: object | None = None
    literature_report: object | None = None
    candidate_literature: object | None = None
    source_review_report: object | None = None
    human_review_result: object | None = None
    report_path: str = ""

    @property
    def has_format_problem(self) -> bool:
        return any(not check.passed for check in self.format_checks)

    @property
    def has_problem(self) -> bool:
        all_checks = (
            self.structure_checks
            + self.file_checks
            + self.trace_checks
            + self.format_checks
        )
        return any(not check.passed for check in all_checks)


def scan_content_markers(markdown: str) -> tuple[list[str], list[str]]:
    missing = sorted(set(re.findall(r"\[待补充：[^\]]+\]", markdown)))
    confirmations = sorted(set(re.findall(r"\[待确认：[^\]]+\]", markdown)))
    return missing, confirmations


def check_required_sections(markdown: str) -> list[FormatCheck]:
    has_title = bool(re.search(r"^#\s+\S+", markdown, re.M))
    checks = [
        FormatCheck(
            "标题",
            has_title,
            "检测到一级标题" if has_title else "缺少一级标题",
        )
    ]
    for section in REQUIRED_SECTIONS:
        present = bool(re.search(rf"^##\s+{re.escape(section)}\s*$", markdown, re.M))
        checks.append(
            FormatCheck(
                section,
                present,
                "检测到章节标题" if present else f"缺少 ## {section}",
            )
        )
    return checks


def check_output_files(paths: dict[str, str]) -> list[FormatCheck]:
    checks: list[FormatCheck] = []
    for label, raw_path in paths.items():
        path = Path(raw_path)
        exists = path.exists() and path.stat().st_size > 0
        checks.append(
            FormatCheck(
                label,
                exists,
                f"{path}，大小 {path.stat().st_size} 字节" if exists else f"{path} 不存在或为空",
            )
        )
    return checks


def check_trace_entries(trace_entries: list) -> list[FormatCheck]:
    checks: list[FormatCheck] = []
    if not trace_entries:
        return [FormatCheck("Trace 记录", False, "未收到 trace_entries")]

    agents = {getattr(entry, "agent", "") for entry in trace_entries}
    expected_agents = {
        "InterviewAgent",
        "PlanningAgent",
        "WritingAgent",
        "WordFormatAgent",
        "VerifierAgent",
    }
    checks.append(
        FormatCheck(
            "Trace Agent 覆盖",
            expected_agents.issubset(agents),
            f"检测到 Agent: {', '.join(sorted(agents))}",
        )
    )

    bad_entries: list[str] = []
    for index, entry in enumerate(trace_entries, 1):
        data = entry if isinstance(entry, dict) else entry.__dict__
        missing = sorted(TRACE_REQUIRED_KEYS - set(data))
        empty_required = [
            key
            for key in TRACE_REQUIRED_KEYS
            if key not in {"failure_reason", "input_keys", "output_keys"}
            and not data.get(key)
        ]
        type_problem = []
        if not isinstance(data.get("input_keys"), list):
            type_problem.append("input_keys")
        if not isinstance(data.get("output_keys"), list):
            type_problem.append("output_keys")
        if missing or empty_required or type_problem:
            bad_entries.append(
                f"#{index} missing={missing or '-'} empty={empty_required or '-'} type={type_problem or '-'}"
            )
    checks.append(
        FormatCheck(
            "Trace 字段完整性",
            not bad_entries,
            "每步均包含 task_id/agent/stage/input_keys/output_keys/tool_call/status/failure_reason/timestamp"
            if not bad_entries
            else "; ".join(bad_entries),
        )
    )
    return checks


def check_docx_format(docx_path: str | Path) -> list[FormatCheck]:
    path = Path(docx_path)
    if not path.exists():
        return [FormatCheck("Word 文档", False, f"不存在 {path}")]

    doc = Document(path)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    heading_count = sum(
        1
        for p in paragraphs
        if p.style is not None and p.style.name.lower().startswith("heading")
    )
    has_body = any(
        p.text.strip()
        and not (p.style and p.style.name.lower().startswith("heading"))
        and p.style.name != "Title"
        for p in paragraphs
    )
    has_reference = any(p.text.strip() in {"参考文献", "References"} for p in paragraphs)
    has_reference_placeholder = any("[待补充：参考文献]" in p.text for p in paragraphs)
    has_reference_section = len(doc.sections) >= 2 or any(
        p.text.strip() in {"参考文献", "References"}
        and bool(p.paragraph_format.page_break_before)
        for p in paragraphs
    )
    has_numbered_reference = any(re.match(r"^\[\d+\]\s+", p.text.strip()) for p in paragraphs)
    has_page_number = any(section.footer.paragraphs for section in doc.sections)
    title = paragraphs[0] if paragraphs else None
    title_centered = bool(title and title.alignment == WD_ALIGN_PARAGRAPH.CENTER)
    title_bold = bool(title and any(run.bold for run in title.runs))
    body_indented = any(
        p.paragraph_format.first_line_indent is not None
        and round(p.paragraph_format.first_line_indent.pt) == 24
        for p in paragraphs
        if not (p.style and p.style.name.lower().startswith(("heading", "title")))
    )

    return [
        FormatCheck("标题居中", title_centered, "首页标题居中" if title_centered else "首页标题未居中"),
        FormatCheck("标题加粗", title_bold, "首页标题加粗" if title_bold else "首页标题未加粗"),
        FormatCheck("标题层级", heading_count > 0, f"Heading 段落数 {heading_count}"),
        FormatCheck(
            "正文统一字体字号",
            has_body,
            "存在正文段落，使用 Normal/正文样式" if has_body else "未发现正文段落",
        ),
        FormatCheck(
            "正文首行缩进",
            body_indented,
            "检测到 24pt 首行缩进" if body_indented else "未检测到正文首行缩进",
        ),
        FormatCheck("页码", has_page_number, "检测到页脚段落" if has_page_number else "未检测到页脚段落"),
        FormatCheck(
            "参考文献单独成节",
            has_reference and has_reference_section,
            "检测到参考文献标题和独立 section"
            if has_reference and has_reference_section
            else "未检测到参考文献独立 section",
        ),
        FormatCheck(
            "参考文献编号清晰",
            has_numbered_reference or not has_reference or has_reference_placeholder,
            "检测到 [n] 参考文献条目" if has_numbered_reference else "无正式参考文献条目，保留待补充占位",
        ),
    ]
