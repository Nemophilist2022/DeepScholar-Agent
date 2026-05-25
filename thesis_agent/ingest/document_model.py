"""DocumentModel — single allowed access point to ``python-docx`` (R2.3, R2.4).

Upper layers (Tools / Evaluators / Diagnoser) must go through this
module instead of importing ``docx`` directly. The model exposes:

- a **read API** returning immutable views (paragraphs / sections / styles
  / tables) — Evaluators and Diagnoser only need this side
- a **controlled write API** (``with dm.write() as w: ...``) — Tools use
  this side; every mutation is recorded in :attr:`last_changes` so the
  Tool can fill ``ToolResult.changed_*`` accurately

Internal note: this is the **only** module in the agent that may import
from ``docx``. T7 ships a lint test that enforces the rule.
"""

from __future__ import annotations

from contextlib import contextmanager
from dataclasses import dataclass, field
from typing import Iterator

from docx import Document
from docx.styles.styles import Styles


# ---------------------------------------------------------------------------
# Read-side immutable views
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ParagraphView:
    """Read-only snapshot of a paragraph at the time of access."""

    index: int
    text: str
    style_name: str


@dataclass(frozen=True)
class SectionView:
    index: int


@dataclass(frozen=True)
class StyleView:
    name: str


@dataclass(frozen=True)
class TableView:
    index: int
    rows: int
    cols: int


# ---------------------------------------------------------------------------
# Change tracking
# ---------------------------------------------------------------------------

@dataclass
class ChangeSet:
    """What changed during the most recent ``with dm.write()`` session."""

    paragraphs: list[int] = field(default_factory=list)
    styles: list[str] = field(default_factory=list)
    sections: list[int] = field(default_factory=list)


# ---------------------------------------------------------------------------
# Writer (yielded inside ``with dm.write()``)
# ---------------------------------------------------------------------------

class DocumentWriter:
    """Receives every mutation; pushes change records into the model."""

    def __init__(self, model: "DocumentModel") -> None:
        self._model = model

    def set_paragraph_text(self, index: int, text: str) -> None:
        para = self._model._doc.paragraphs[index]
        # Preserve the first run's formatting if any; otherwise just
        # replace text. This mirrors thesis_formatter.numbering._set_para_text
        # but kept minimal for MVP.
        if para.runs:
            para.runs[0].text = text
            for r in para.runs[1:]:
                r.text = ""
        else:
            para.add_run(text)
        self._model._record_paragraph(index)

    def mark_style_dirty(self, style_name: str) -> None:
        """Tools that mutate a style via the underlying python-docx API
        should call this so the change is tracked."""
        self._model._record_style(style_name)

    def mark_section_dirty(self, index: int) -> None:
        self._model._record_section(index)

    @property
    def raw(self):
        """Escape hatch for Tools that wrap legacy ``thesis_formatter``
        functions which already mutate the underlying docx. Should be
        paired with explicit ``mark_*_dirty`` calls.
        """
        return self._model._doc


# ---------------------------------------------------------------------------
# Top-level model
# ---------------------------------------------------------------------------

class DocumentModel:
    """The single authority on the in-memory docx state."""

    def __init__(self, docx_doc) -> None:
        self._doc = docx_doc
        self._last_changes: ChangeSet = ChangeSet()

    # ----- Construction -----

    @classmethod
    def from_path(cls, path: str) -> "DocumentModel":
        return cls(Document(path))

    # ----- Read API -----

    def paragraphs(self) -> tuple[ParagraphView, ...]:
        return tuple(
            ParagraphView(
                index=i,
                text=p.text,
                style_name=(p.style.name if p.style is not None else ""),
            )
            for i, p in enumerate(self._doc.paragraphs)
        )

    def sections(self) -> tuple[SectionView, ...]:
        return tuple(SectionView(index=i) for i in range(len(self._doc.sections)))

    def styles(self) -> tuple[StyleView, ...]:
        styles_obj: Styles = self._doc.styles
        return tuple(StyleView(name=s.name) for s in styles_obj)

    def tables(self) -> tuple[TableView, ...]:
        out = []
        for i, t in enumerate(self._doc.tables):
            rows = len(t.rows)
            cols = len(t.columns) if t.columns else (len(t.rows[0].cells) if rows else 0)
            out.append(TableView(index=i, rows=rows, cols=cols))
        return tuple(out)

    # ----- Controlled write API -----

    @contextmanager
    def write(self) -> Iterator[DocumentWriter]:
        """Open a write session. ``last_changes`` is reset on entry and
        accumulated until the ``with`` block exits successfully."""
        self._last_changes = ChangeSet()
        writer = DocumentWriter(self)
        yield writer

    @property
    def last_changes(self) -> ChangeSet:
        return self._last_changes

    def _record_paragraph(self, index: int) -> None:
        if index not in self._last_changes.paragraphs:
            self._last_changes.paragraphs.append(index)

    def _record_style(self, style_name: str) -> None:
        if style_name not in self._last_changes.styles:
            self._last_changes.styles.append(style_name)

    def _record_section(self, index: int) -> None:
        if index not in self._last_changes.sections:
            self._last_changes.sections.append(index)

    # ----- Persistence -----

    def save(self, path: str) -> None:
        self._doc.save(path)
