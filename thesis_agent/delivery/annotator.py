"""Highlight partial / failed paragraphs in the output docx (R7.6, D8).

This is the **fallback** annotator from D8: a light-blue paragraph
shading (``#E6F3FF``) plus a short trailing inline note listing the
rule ids that flagged the paragraph. The colour was chosen to differ
from Word's default yellow highlight so users can tell agent
annotations apart from their own.

Word Comment XML injection (the preferred D8 form) is deferred to a
later milestone — see the ``risks`` section in the B2 wrap-up.

Public API:
    annotate(docx_path, delivery)

Failure policy:
    Annotation is best-effort. Any IO / OOXML error is caught here so
    the harness can carry on; we just record a ``warnings`` list.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import Any, Iterable, Optional

# python-docx is allowed in the delivery layer because we are writing
# the output docx, not reading the user's content.
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# The fixed shading colour (D8 fallback choice).
ANNOTATION_FILL = "E6F3FF"
ANNOTATION_NOTE_PREFIX = "  ⚠ "


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _delivery_items(delivery) -> list:
    """Pull the iterable of items out of a :class:`DeliveryReport`."""
    return list(getattr(delivery, "items", []) or [])


def _annotatable(item) -> bool:
    """Return True if *item* is a kind we should mark in docx."""
    return getattr(item, "status", None) in ("partial", "failed")


def _locator_index(locator: dict) -> Optional[int]:
    """If the locator pins a specific paragraph, return its index."""
    if not isinstance(locator, dict):
        return None
    val = locator.get("paragraph_index")
    if isinstance(val, int):
        return val
    return None


def _locator_style_name(locator: dict) -> Optional[str]:
    """Some rules locate via a style name (e.g. Heading 1). Return it."""
    if not isinstance(locator, dict):
        return None
    name = locator.get("style_name")
    if isinstance(name, str) and name:
        return name
    return None


def _locator_front_matter(locator: dict) -> Optional[str]:
    if not isinstance(locator, dict):
        return None
    target = locator.get("front_matter")
    if isinstance(target, str) and target:
        return target
    return None


def _set_paragraph_shading(para, fill: str) -> None:
    """Apply ``<w:shd w:fill="<fill>"/>`` to the paragraph's pPr."""
    p_pr = para._element.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _append_inline_note(para, note: str) -> None:
    """Append a small trailing run with the rule ids."""
    run = para.add_run(note)
    run.italic = True
    # Subtle grey so it visually differs from body text without
    # competing with the blue background.
    rgb = run.font.color
    rgb.rgb = None  # let docx use the default style colour
    # The italic style is enough for visual differentiation across
    # Word versions that ignore custom RGB on a freshly-added run.


# ---------------------------------------------------------------------------
# Locator resolution
# ---------------------------------------------------------------------------

@dataclass
class _ParagraphTarget:
    """A resolved paragraph object plus the rule ids attached to it."""

    paragraph: Any
    rule_ids: list[str]


def _resolve_targets(doc, items: list) -> list[_ParagraphTarget]:
    """Group annotatable items by the paragraph they should mark.

    Items whose locator can't pin a single paragraph are dropped (they
    still appear in the report — this is purely the visual side).
    """
    by_index: dict[int, list[str]] = {}
    by_style: dict[str, list[str]] = {}
    by_front: dict[str, list[str]] = {}

    for it in items:
        if not _annotatable(it):
            continue
        loc = getattr(it, "locator", {}) or {}
        rule_id = getattr(it, "rule_id", "")
        idx = _locator_index(loc)
        if idx is not None:
            by_index.setdefault(idx, []).append(rule_id)
            continue
        style = _locator_style_name(loc)
        if style and style != "Normal":
            # Whole-document Normal-style violations are too broad to
            # highlight every paragraph. Heading{1..4} etc. are useful.
            by_style.setdefault(style, []).append(rule_id)
            continue
        fm = _locator_front_matter(loc)
        if fm:
            by_front.setdefault(fm, []).append(rule_id)
            continue
        # Other locator shapes (all_sections / all_tables / caption /
        # toc / page_numbers) are document-wide — no single paragraph
        # to highlight.

    targets: list[_ParagraphTarget] = []

    # paragraph_index — direct
    for idx, rule_ids in by_index.items():
        if 0 <= idx < len(doc.paragraphs):
            targets.append(_ParagraphTarget(
                paragraph=doc.paragraphs[idx],
                rule_ids=rule_ids,
            ))

    # style_name — every paragraph with that style
    for style_name, rule_ids in by_style.items():
        for para in doc.paragraphs:
            if para.style and para.style.name == style_name:
                targets.append(_ParagraphTarget(
                    paragraph=para, rule_ids=list(rule_ids)
                ))

    # front_matter — match by text prefix / equality, mirrors the
    # heuristics in evaluators.checks.check_front_matter.
    for fm, rule_ids in by_front.items():
        para = _find_front_matter_paragraph(doc, fm)
        if para is not None:
            targets.append(_ParagraphTarget(
                paragraph=para, rule_ids=list(rule_ids)
            ))

    return targets


def _find_front_matter_paragraph(doc, target: str):
    """Best-effort locator: find the first paragraph that looks like
    the requested front-matter element."""
    import re

    def _match(text: str) -> bool:
        compact = re.sub(r"\s+", "", text or "").replace("\u3000", "")
        if target == "cn_abstract":
            return compact == "摘要"
        if target == "cn_keywords":
            return bool(re.match(r"^\s*关键词\s*[：:]", text or ""))
        if target == "en_abstract":
            return (
                bool(re.match(r"(?i)^\s*Abstract\s*[:：]", text or ""))
                or bool(re.match(r"(?i)^\s*Abstract\s*$", text or ""))
            )
        if target == "en_keywords":
            return bool(re.match(r"(?i)^\s*Key\s*words?\s*[:：]",
                                 re.sub(r"\s+", " ", text or "")))
        return False

    for para in doc.paragraphs:
        if _match(para.text or ""):
            return para
    return None


def _merge_targets_per_paragraph(
    targets: Iterable[_ParagraphTarget],
) -> list[_ParagraphTarget]:
    """Same paragraph from multiple locator shapes → merge rule ids."""
    seen: dict[int, _ParagraphTarget] = {}
    for t in targets:
        key = id(t.paragraph._element)
        if key in seen:
            existing = seen[key]
            for rid in t.rule_ids:
                if rid not in existing.rule_ids:
                    existing.rule_ids.append(rid)
        else:
            seen[key] = _ParagraphTarget(
                paragraph=t.paragraph,
                rule_ids=list(t.rule_ids),
            )
    return list(seen.values())


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

@dataclass
class AnnotationResult:
    annotated_paragraphs: int = 0
    skipped_items: int = 0
    warnings: list[str] = None  # type: ignore[assignment]

    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []


def annotate(docx_path: str, delivery) -> AnnotationResult:
    """Apply paragraph highlights for every partial / failed item that
    can be pinned to a specific paragraph.

    Always saves the docx in place. Returns a small summary so the
    harness can write the count into the trace.
    """
    result = AnnotationResult()
    items = _delivery_items(delivery)
    if not items:
        return result

    annotatable = [it for it in items if _annotatable(it)]
    if not annotatable:
        return result

    try:
        doc = Document(docx_path)
    except Exception as exc:  # IO / format error — skip annotation
        result.warnings.append(f"open_failed: {exc}")
        return result

    try:
        targets = _resolve_targets(doc, annotatable)
        targets = _merge_targets_per_paragraph(targets)
        for t in targets:
            try:
                _set_paragraph_shading(t.paragraph, ANNOTATION_FILL)
                note = ANNOTATION_NOTE_PREFIX + ", ".join(t.rule_ids)
                _append_inline_note(t.paragraph, note)
                result.annotated_paragraphs += 1
            except Exception as exc:
                result.warnings.append(
                    f"shade_failed rule_ids={t.rule_ids}: {exc}"
                )
        result.skipped_items = max(0, len(annotatable) - len(targets))
        doc.save(docx_path)
    except Exception as exc:
        result.warnings.append(f"annotate_failed: {exc}")
    return result
