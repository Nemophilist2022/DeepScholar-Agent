import copy
import os
import re
import shutil
import sys
import tempfile
import zipfile
import subprocess

from thesis_config import resolve_config
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ._common import (
    _ALIGN_MAP, set_style_font, set_run_font, set_para_runs_font,
    set_table_border, clear_table_border,
    _ensure_keep_next, _set_para_spacing, _check_caption_numbering,
    is_heading, contains_cjk, normalize_cn_keywords, normalize_en_keywords,
    get_paragraph_heading_level, get_heading_style, is_heading_style,
    _ALL_HEADING_NAMES, parse_length, apply_line_spacing, apply_paragraph_spacing,
)
from ._titles import _find_special_display, _get_special_title_map, _detect_front_matter
from .headings import auto_assign_heading_styles, renumber_headings, normalize_heading_spacing, demote_abstract_heading_styles
from .references import check_citations, apply_ref_crosslinks
from .page import normalize_sections, setup_page_numbers, setup_page_numbers_strict, insert_page_break_after, find_first_body_heading
from .headers import setup_headers
from .toc import insert_toc, ensure_toc_styles
from .cover import _has_cover, find_existing_cover_end, insert_cover_and_declaration
from .structure import validate_structure


def _build_insert_cover_vbs():
    # 以原文件为基础，用剪贴板粘贴封面（PasteAndFormat 保留源格式），再加分节符。
    return """Option Explicit
Const wdStory = 6
Const wdSectionBreakNextPage = 7
Const wdFormatOriginalFormatting = 16
Dim objWord, targetDoc, coverDoc, args, targetPath, coverPath
Set args = WScript.Arguments
If args.Count < 2 Then WScript.Quit 1
targetPath = args(0): coverPath = args(1)

On Error Resume Next
Set objWord = CreateObject("Word.Application")
If Err.Number <> 0 Then WScript.Quit 1
On Error GoTo 0

objWord.Visible = False: objWord.DisplayAlerts = 0

' 1. Open cover, copy all content to clipboard
On Error Resume Next
Set coverDoc = objWord.Documents.Open(coverPath)
If Err.Number <> 0 Then
    objWord.Quit
    WScript.Quit 1
End If
On Error GoTo 0
coverDoc.Content.Copy
coverDoc.Close False

' 2. Open target, paste cover at start with original formatting
On Error Resume Next
Set targetDoc = objWord.Documents.Open(targetPath)
If Err.Number <> 0 Then
    objWord.Quit
    WScript.Quit 1
End If
On Error GoTo 0

objWord.Selection.HomeKey wdStory
objWord.Selection.PasteAndFormat wdFormatOriginalFormatting

' 3. After paste, Selection is at end of cover content - insert section break
objWord.Selection.InsertBreak wdSectionBreakNextPage

' 4. Clear cover section headers/footers
ClearSectionHeaderFooter targetDoc.Sections(1)
If targetDoc.Sections.Count >= 2 Then
    UnlinkSectionHeaderFooter targetDoc.Sections(2)
End If

targetDoc.Save
targetDoc.Close False
objWord.Quit

Set targetDoc = Nothing
Set coverDoc = Nothing
Set objWord = Nothing
WScript.Echo "Done"

Sub ClearSectionHeaderFooter(section)
    Dim idx
    On Error Resume Next
    For idx = 1 To 3
        section.Headers(idx).LinkToPrevious = False
        section.Headers(idx).Range.Text = ""
        section.Footers(idx).LinkToPrevious = False
        section.Footers(idx).Range.Text = ""
    Next
    Err.Clear
    On Error GoTo 0
End Sub

Sub UnlinkSectionHeaderFooter(section)
    Dim idx
    On Error Resume Next
    For idx = 1 To 3
        section.Headers(idx).LinkToPrevious = False
        section.Footers(idx).LinkToPrevious = False
    Next
    Err.Clear
    On Error GoTo 0
End Sub
"""
def _insert_cover_via_vbs(target_path, cover_path):
    """使用嵌入的 VBS 代码插入封面，保留完整格式."""
    import tempfile

    vbs_code = _build_insert_cover_vbs()
    vbs_path = None

    try:
        with tempfile.NamedTemporaryFile(mode="w", suffix=".vbs", delete=False) as f:
            f.write(vbs_code)
            vbs_path = f.name

        result = subprocess.run(
            ["cscript", "//NoLogo", vbs_path, target_path, cover_path],
            capture_output=True,
            text=True,
            encoding="gbk",
            errors="replace",
            timeout=60
        )

        return result.returncode == 0, result.stderr if result.returncode != 0 else ""
    except Exception as e:
        return False, str(e)
    finally:
        if vbs_path:
            try:
                os.unlink(vbs_path)
            except OSError:
                pass

def _insert_space_at_offset(para, offset):
    """Insert a single space character at *offset* within the paragraph text,
    preserving all existing run formatting."""
    pos = 0
    for run in para.runs:
        rt = run.text or ""
        run_end = pos + len(rt)
        if pos <= offset < run_end:
            local = offset - pos
            run.text = rt[:local] + " " + rt[local:]
            return
        pos = run_end


def _format_tables(doc, cfg):
    """Format all tables: cell alignment, fonts, borders (three-line table style)."""
    latin = cfg["fonts"]["latin"]
    body_font = cfg["fonts"]["body"]
    caption_size = parse_length(cfg["sizes"]["caption"])
    tbl_cfg = cfg["table"]
    tbl_cell_align = _ALIGN_MAP.get(tbl_cfg.get("cell_align", "center"))
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'autofit')

        rows = len(table.rows)
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    if tbl_cell_align is not None:
                        p.alignment = tbl_cell_align
                    apply_paragraph_spacing(p.paragraph_format, "before", 0)
                    apply_paragraph_spacing(p.paragraph_format, "after", 0)
                    p.paragraph_format.first_line_indent = parse_length(0)
                    apply_line_spacing(p.paragraph_format, tbl_cfg["line_spacing"])
                    set_para_runs_font(p, east_asia=body_font, size_pt=caption_size,
                                       bold=False, latin=latin)

                clear_table_border(cell, "left")
                clear_table_border(cell, "right")
                clear_table_border(cell, "insideV")
                clear_table_border(cell, "insideH")

                if r_idx == 0:
                    set_table_border(cell, "top", sz=tbl_cfg["top_border_sz"])
                    set_table_border(cell, "bottom", sz=tbl_cfg["header_border_sz"])
                if r_idx == rows - 1:
                    set_table_border(cell, "bottom", sz=tbl_cfg["bottom_border_sz"])


def _format_captions(doc, cfg, preserved_para_ids):
    """Format figure/table/subfigure/note/source captions and apply keep-with-next."""
    latin = cfg["fonts"]["latin"]
    body_font = cfg["fonts"]["body"]
    caption_size = parse_length(cfg["sizes"]["caption"])
    note_size = parse_length(cfg["sizes"]["note"])
    body_ls = cfg["body"]["line_spacing"]
    spacing_line = parse_length(cfg["sizes"]["body"])

    cap_cfg = cfg.get("captions", {})
    fig_pat = cap_cfg.get("figure_pattern", r"^图\s*\d")
    tbl_pat = cap_cfg.get("table_pattern", r"^(续)?表\s*\d")
    subfig_pat = cap_cfg.get("subfigure_pattern", r"^\([a-z]\)")
    note_pat = cap_cfg.get("note_pattern", r"^注[：:]")
    source_pat = r"^(资料)?来源\s*[：:]"
    cap_ls = cap_cfg.get("line_spacing", body_ls)

    def _is_preserved(para):
        return id(para._element) in preserved_para_ids

    _cap_space_re = re.compile(r"^((?:图|表|Figure|Table)\s*[A-Z]?\d+)(\S)", re.I)
    for para in doc.paragraphs:
        if _is_preserved(para):
            continue
        has_seq = any("instrText" in str(run._element) for run in para.runs)
        if has_seq:
            continue
        t = para.text.strip()
        m = _cap_space_re.match(t)
        if m:
            _insert_space_at_offset(para, len(m.group(1)))

    for para in doc.paragraphs:
        if _is_preserved(para):
            continue
        t = para.text.strip()
        if re.match(fig_pat, t) or re.match(r"^Figure\s*\d", t, re.I) or re.match(r"^图[A-Z]\d+", t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = parse_length(0)
            apply_line_spacing(para.paragraph_format, cap_ls)
            para.paragraph_format.space_after = spacing_line
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(tbl_pat, t) or re.match(r"^Table\s*\d", t, re.I) or re.match(r"^(续)?表[A-Z]\d+", t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = parse_length(0)
            apply_line_spacing(para.paragraph_format, cap_ls)
            para.paragraph_format.space_before = spacing_line
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(subfig_pat, t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = parse_length(0)
            apply_line_spacing(para.paragraph_format, cap_ls)
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(note_pat, t) or re.match(source_pat, t):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = parse_length(0)
            apply_line_spacing(para.paragraph_format, cap_ls)
            set_para_runs_font(para, east_asia=body_font, size_pt=note_size,
                               bold=False, latin=latin)

    if cap_cfg.get("keep_with_next", True):
        body_el = doc.element.body
        children = list(body_el)
        for i, el in enumerate(children):
            if el.tag == qn("w:tbl"):
                if i > 0 and children[i - 1].tag == qn("w:p"):
                    prev_text = "".join(
                        (nd.text or "") for nd in children[i - 1].iter(qn("w:t"))
                    ).strip()
                    if re.match(tbl_pat, prev_text) or re.match(r"^Table\s*\d", prev_text, re.I) or re.match(r"^(续)?表[A-Z]\d+", prev_text):
                        _ensure_keep_next(children[i - 1])
                if i + 1 < len(children) and children[i + 1].tag == qn("w:p"):
                    nt = "".join(
                        (nd.text or "") for nd in children[i + 1].iter(qn("w:t"))
                    ).strip()
                    if nt and not re.match(note_pat, nt) and not re.match(source_pat, nt):
                        _set_para_spacing(children[i + 1], "before", spacing_line)
            elif el.tag == qn("w:p") and el.findall(".//" + qn("w:drawing")):
                if i + 1 < len(children) and children[i + 1].tag == qn("w:p"):
                    next_text = "".join(
                        (nd.text or "") for nd in children[i + 1].iter(qn("w:t"))
                    ).strip()
                    if re.match(fig_pat, next_text) or re.match(r"^Figure\s*\d", next_text, re.I) or re.match(r"^图[A-Z]\d+", next_text):
                        _ensure_keep_next(el)
                if i > 0 and children[i - 1].tag == qn("w:p"):
                    pt = "".join(
                        (nd.text or "") for nd in children[i - 1].iter(qn("w:t"))
                    ).strip()
                    if pt and not re.match(fig_pat, pt) and not re.match(subfig_pat, pt):
                        _set_para_spacing(el, "before", spacing_line)

    return cap_cfg, fig_pat, tbl_pat


def apply_format(input_path, output_path, config=None, config_path=None):
    if config is None:
        config, config_path = resolve_config(input_path=input_path)
    cfg = config

    latin = cfg["fonts"]["latin"]
    body_font = cfg["fonts"]["body"]
    body_size = parse_length(cfg["sizes"]["body"])
    body_ls = cfg["body"]["line_spacing"]
    body_indent = parse_length(cfg["body"]["first_line_indent"])
    body_align = _ALIGN_MAP.get(cfg["body"]["align"])

    h1_font = cfg["fonts"]["h1"]
    h1_size = parse_length(cfg["sizes"]["h1"])
    h2_font = cfg["fonts"]["h2"]
    h2_size = parse_length(cfg["sizes"]["h2"])
    h3_font = cfg["fonts"]["h3"]
    h3_size = parse_length(cfg["sizes"]["h3"])
    h4_font = cfg["fonts"]["h4"]
    h4_size = parse_length(cfg["sizes"]["h4"])

    def _bold_val(v):
        return None if v == "keep" else v

    h1_bold = _bold_val(cfg["headings"]["h1"]["bold"])
    h1_align = _ALIGN_MAP.get(cfg["headings"]["h1"]["align"])
    h2_bold = _bold_val(cfg["headings"]["h2"]["bold"])
    h2_align = _ALIGN_MAP.get(cfg["headings"]["h2"]["align"])
    h3_bold = _bold_val(cfg["headings"]["h3"]["bold"])
    h3_align = _ALIGN_MAP.get(cfg["headings"]["h3"]["align"])
    h4_bold = _bold_val(cfg["headings"]["h4"]["bold"])
    h4_align = _ALIGN_MAP.get(cfg["headings"]["h4"]["align"])

    fn_size = parse_length(cfg["sizes"]["footnote"])

    st_map = _get_special_title_map(cfg)
    sec = cfg["sections"]
    ref_key = "参考文献"
    toc_key = "目录"
    toc_cfg = cfg.get("toc", {})
    toc_enabled = toc_cfg.get("enabled", True)
    toc_only = toc_cfg.get("only_insert", False)
    cover_cfg = cfg.get("cover", {})
    cover_only = cover_cfg.get("only_insert", False)
    page_numbers_cfg = cfg.get("page_numbers", {})
    page_numbers_only = page_numbers_cfg.get("only_insert", False)
    header_footer_cfg = cfg.get("header_footer", {})
    header_only = header_footer_cfg.get("only_insert", False)
    custom_cover = cover_cfg.get("custom_docx", "")
    use_custom_cover = bool(custom_cover and os.path.isfile(custom_cover))

    local_modes = {
        "toc": toc_only,
        "cover": cover_only,
        "page_numbers": page_numbers_only,
        "header_footer": header_only,
    }
    active_local_modes = [name for name, enabled in local_modes.items() if enabled]
    if len(active_local_modes) > 1:
        raise RuntimeError("单独处理模式不能同时启用多个。")

    if cover_only:
        if not use_custom_cover:
            raise RuntimeError("仅插入外部封面模式需要提供有效的自定义封面 .docx")
        shutil.copy2(input_path, output_path)
        success, err = _insert_cover_via_vbs(output_path, custom_cover)
        if not success:
            raise RuntimeError(f"自定义封面插入失败: {err}")
        cfg.setdefault("_runtime", {})["custom_cover_sections"] = 1
        cfg["_runtime"]["cover_only"] = True
        cfg["_runtime"]["local_mode"] = "cover"
        return []

    doc = Document(input_path)

    if page_numbers_only:
        setup_page_numbers_strict(doc, cfg)
        cfg.setdefault("_runtime", {})["local_mode"] = "page_numbers"
        doc.save(output_path)
        return []

    if header_only:
        setup_headers(doc, cfg)
        cfg.setdefault("_runtime", {})["local_mode"] = "header_footer"
        doc.save(output_path)
        return []
    def _configure_toc_styles():
        if "TOC Heading" in doc.styles:
            st = doc.styles["TOC Heading"]
            toc_h_font = cfg["toc"].get("h1_font", h1_font)
            toc_h_size = parse_length(cfg["toc"].get("h1_font_size", cfg["sizes"]["h1"]))
            set_style_font(st, east_asia=toc_h_font, size_pt=toc_h_size, bold=True, latin=latin)
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _toc_h_sb = cfg["toc"].get("space_before", 0)
            _toc_h_sa = cfg["toc"].get("space_after", 0)
            apply_paragraph_spacing(st.paragraph_format, "before", _toc_h_sb)
            apply_paragraph_spacing(st.paragraph_format, "after", _toc_h_sa)
        ensure_toc_styles(doc, cfg)


    if toc_only:
        if toc_enabled:
            from .headings import assign_outline_levels_for_toc
            assign_outline_levels_for_toc(doc, cfg)
            _configure_toc_styles()
            insert_toc(doc, cfg)
        cfg.setdefault("_runtime", {})["local_mode"] = "toc"
        doc.save(output_path)
        return []
    fm_mode = cfg.get("front_matter", {}).get("mode", "auto")

    preserve_until_idx = 0
    first_body_heading = find_first_body_heading(doc, cfg)
    if fm_mode == "skip":
        if first_body_heading is None:
            preserve_until_idx = len(doc.paragraphs)
        else:
            for idx, para in enumerate(doc.paragraphs):
                if para._element is first_body_heading._element:
                    preserve_until_idx = idx
                    break
    elif not cfg.get("cover", {}).get("enabled", True):
        preserve_until_idx = find_existing_cover_end(doc, cfg)

    preserved_para_ids = {id(para._element) for para in doc.paragraphs[:preserve_until_idx]}
    preserve_front_matter = bool(preserved_para_ids)
    preserved_front_snapshots = {
        id(para._element): copy.deepcopy(para._element)
        for para in doc.paragraphs[:preserve_until_idx]
    } if preserve_front_matter else {}

    def _is_preserved_front_para(para):
        return id(para._element) in preserved_para_ids

    def _restore_preserved_front_paragraphs():
        if not preserved_front_snapshots:
            return
        for para in doc.paragraphs:
            el_id = id(para._element)
            if el_id not in preserved_front_snapshots:
                continue

            current_el = para._element
            parent = current_el.getparent()
            if parent is None:
                continue

            original_el = preserved_front_snapshots[el_id]

            current_ppr = current_el.find(qn("w:pPr"))
            current_sectpr = None
            if current_ppr is not None:
                sectpr = current_ppr.find(qn("w:sectPr"))
                if sectpr is not None:
                    current_sectpr = copy.deepcopy(sectpr)

            restored_el = copy.deepcopy(original_el)
            if current_sectpr is not None:
                restored_ppr = restored_el.find(qn("w:pPr"))
                if restored_ppr is None:
                    restored_ppr = OxmlElement("w:pPr")
                    restored_el.insert(0, restored_ppr)
                for old_sectpr in list(restored_ppr.findall(qn("w:sectPr"))):
                    restored_ppr.remove(old_sectpr)
                restored_ppr.append(current_sectpr)

            parent.replace(current_el, restored_el)

    cover_end_idx = preserve_until_idx

    warnings = []
    from . import numbering
    requested_caption_mode, raw_caption_mode, raw_caption_mode_warnings, raw_caption_reasons = numbering.resolve_caption_mode(doc, cfg)
    cfg.setdefault("_runtime", {})
    cfg["_runtime"]["caption_mode_requested"] = requested_caption_mode
    cfg["_runtime"]["caption_mode_raw_effective"] = raw_caption_mode
    cfg["_runtime"]["caption_mode_raw_precheck_passed"] = requested_caption_mode != numbering.CAPTION_MODE_DYNAMIC or not raw_caption_reasons

    auto_changes = auto_assign_heading_styles(doc, cfg, skip_para_ids=preserved_para_ids)
    if auto_changes:
        print(f"自动识别标题 ({len(auto_changes)} 个):", file=sys.stderr)
        for c in auto_changes:
            print(c, file=sys.stderr)

    try:
        warnings.extend(validate_structure(doc, cfg) or [])
    except Exception as exc:
        print(f"结构检查出错（已跳过，继续排版）: {exc}", file=sys.stderr)
    normalize_sections(doc, cfg)

    renum_changes = []
    if sec.get("renumber_headings", False):
        renum_changes = renumber_headings(doc, cfg, skip_para_ids=preserved_para_ids)

    requested_caption_mode, effective_caption_mode, caption_mode_warnings = numbering.resolve_caption_mode_after_normalization(
        doc, cfg, raw_reasons=raw_caption_reasons
    )
    cfg["_runtime"]["caption_mode_effective"] = effective_caption_mode
    cfg["_runtime"]["caption_mode_precheck_passed"] = effective_caption_mode == numbering.CAPTION_MODE_DYNAMIC
    warnings.extend(raw_caption_mode_warnings)
    warnings.extend(caption_mode_warnings)

    normalize_heading_spacing(doc, cfg, skip_para_ids=preserved_para_ids)

    if not preserve_front_matter:
        for style_name in ["Normal", "Body Text", "First Paragraph", "_Style 2"]:
            if style_name in doc.styles:
                set_style_font(doc.styles[style_name], east_asia=body_font,
                               size_pt=body_size, bold=False, latin=latin)

    def _set_heading_style(level, font, size, bold, align, hcfg):
        style = get_heading_style(doc, level)
        if style is None:
            return
        set_style_font(style, east_asia=font, size_pt=size, bold=bold, latin=latin)
        if align is not None:
            style.paragraph_format.alignment = align
        sb = hcfg.get("space_before", 0)
        sa = hcfg.get("space_after", 0)
        if sb >= 0:
            apply_paragraph_spacing(style.paragraph_format, "before", sb)
        if sa >= 0:
            apply_paragraph_spacing(style.paragraph_format, "after", sa)

    if not preserve_front_matter:
        _set_heading_style(1, h1_font, h1_size, h1_bold, h1_align, cfg["headings"]["h1"])
        _set_heading_style(2, h2_font, h2_size, h2_bold, h2_align, cfg["headings"]["h2"])
        _set_heading_style(3, h3_font, h3_size, h3_bold, h3_align, cfg["headings"]["h3"])
        _set_heading_style(4, h4_font, h4_size, h4_bold, h4_align, cfg["headings"]["h4"])

    _configure_toc_styles()

    toc_content_font = cfg["toc"].get("font", body_font)
    toc_content_size = parse_length(cfg["toc"].get("font_size", cfg["sizes"]["body"]))
    toc_content_bold = cfg["toc"].get("bold", False)
    toc_h1_font = cfg["toc"].get("h1_font", cfg["fonts"]["h1"])
    toc_h1_size = parse_length(cfg["toc"].get("h1_font_size", cfg["sizes"]["h1"]))
    toc_h1_bold = cfg["toc"].get("h1_bold", False)
    toc_content_ls = cfg["toc"].get("line_spacing", body_ls)
    toc_sb = cfg["toc"].get("space_before", 0)
    toc_sa = cfg["toc"].get("space_after", 0)

    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        sn = para.style.name if para.style else ""
        if sn.lower().startswith("toc ") or sn == "样式3":
            is_toc1 = sn.lower() == "toc 1"
            para.paragraph_format.first_line_indent = parse_length(0)
            apply_line_spacing(para.paragraph_format, toc_content_ls)
            apply_paragraph_spacing(para.paragraph_format, "before", toc_sb)
            apply_paragraph_spacing(para.paragraph_format, "after", toc_sa)
            ea = toc_h1_font if is_toc1 else toc_content_font
            sz = toc_h1_size if is_toc1 else toc_content_size
            bold = toc_h1_bold if is_toc1 else toc_content_bold
            set_para_runs_font(para, east_asia=ea, size_pt=sz,
                               bold=bold, latin=latin)

    for name in ["Footnote Text", "Footnote Reference"]:
        if name in doc.styles:
            set_style_font(doc.styles[name], east_asia=body_font, size_pt=fn_size,
                           bold=False, latin=latin)
    if "Footnote Text" in doc.styles:
        ft = doc.styles["Footnote Text"]
        apply_line_spacing(ft.paragraph_format, cfg["footnote"]["line_spacing"])
        ft.paragraph_format.first_line_indent = parse_length(0)
        _fn_align = _ALIGN_MAP.get(cfg["footnote"].get("align", "justify"))
        if _fn_align is not None:
            ft.paragraph_format.alignment = _fn_align

    for name in ["Hyperlink", "超链接"]:
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.color.rgb = RGBColor(0, 0, 0)
            st.font.underline = False

    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        level = get_paragraph_heading_level(para)
        sn = para.style.name if para.style else ""
        if sn.lower().startswith("toc ") or sn == "样式3":
            continue
        if level is not None:
            hkey = {1: "h1", 2: "h2", 3: "h3", 4: "h4"}.get(level, "h1")
            hcfg = cfg["headings"].get(hkey, {})
            if hcfg.get("space_before", 0) >= 0:
                apply_paragraph_spacing(para.paragraph_format, "before", 0)
            if hcfg.get("space_after", 0) >= 0:
                apply_paragraph_spacing(para.paragraph_format, "after", 0)
        else:
            apply_paragraph_spacing(para.paragraph_format, "before", cfg["body"].get("space_before", 0))
            apply_paragraph_spacing(para.paragraph_format, "after", cfg["body"].get("space_after", 0))
        pf = para.paragraph_format
        if para.style and para.style.name in ["Normal", "Body Text", "First Paragraph", "_Style 2"]:
            if body_align is not None:
                pf.alignment = body_align
            pf.first_line_indent = body_indent
            apply_line_spacing(pf, body_ls)
            set_para_runs_font(para, east_asia=body_font, size_pt=body_size,
                               bold=False, latin=latin)

    fm_mode = cfg.get("front_matter", {}).get("mode", "auto")
    has_fm = (fm_mode == "format") or \
             (fm_mode == "auto" and _detect_front_matter(doc, cfg))

    cn_kw_para = None
    en_kw_para = None

    if has_fm:
        first_body_h1 = find_first_body_heading(doc, cfg)
        if first_body_h1 is None:
            first_h1_idx = len(doc.paragraphs)
        else:
            first_h1_idx = next(
                (i for i, para in enumerate(doc.paragraphs) if para._p is first_body_h1._p),
                len(doc.paragraphs),
            )

        front = doc.paragraphs[cover_end_idx:first_h1_idx]
        non_empty = [p for p in front if p.text.strip()]

        cn_kw_re = sec.get("cn_keywords_pattern", r"^\s*关键词\s*[：:]")
        en_abs_re = sec.get("en_abstract_pattern", r"(?i)^\s*Abstract\s*[\uff1a:]")
        en_kw_re = sec.get("en_keywords_pattern", r"(?i)^\s*Key\s*words\s*[\uff1a:]")
        abstract_display = _find_special_display(cfg, "摘要")
        abstract_display_key = abstract_display.replace(" ", "").replace("\u3000", "")
        front_has_cjk = any(contains_cjk(p.text.strip()) for p in non_empty)
        front_has_english = any(
            bool(re.search(r"[A-Za-z]", p.text.strip())) and not contains_cjk(p.text.strip())
            for p in non_empty
        )
        has_explicit_cn_abstract = any(
            p.text.strip().replace(" ", "").replace("　", "") in {"摘要", abstract_display_key}
            for p in non_empty
        )
        has_explicit_en_abstract = any(
            re.match(en_abs_re, p.text.strip()) or re.match(r"(?i)^\s*Abstract\s*$", p.text.strip())
            for p in non_empty
        )
        if non_empty and front_has_cjk and not has_explicit_cn_abstract:
            warnings.append("前置页缺少明确的摘要/Abstract标题，已保留原文，未强行补写。")

        past_abstract = False
        en_title_seen = False
        normal_style = doc.styles["Normal"] if "Normal" in doc.styles else None

        for idx, p in enumerate(non_empty):
            t = p.text.strip()
            t_nospace = t.replace(" ", "").replace("\u3000", "")
            if idx == 0 and t_nospace in {"摘要", abstract_display_key}:
                if normal_style is not None:
                    p.style = normal_style
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = parse_length(0)
                apply_line_spacing(p.paragraph_format, body_ls)
                r = p.add_run(abstract_display)
                set_run_font(r, east_asia=h1_font, size_pt=h1_size, bold=True, latin=latin)
            elif re.match(cn_kw_re, t):
                cn_kw_para = p
                normalized = normalize_cn_keywords(t) or t
                content = normalized.split("：", 1)[1] if "：" in normalized else ""
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = parse_length(0)
                apply_line_spacing(p.paragraph_format, body_ls)
                r1 = p.add_run("关键词：")
                set_run_font(r1, east_asia=h1_font, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=body_font, size_pt=body_size, bold=False, latin=latin)
            elif re.match(r"(?i)^\s*Abstract\s*$", t):
                past_abstract = True
                if normal_style is not None:
                    p.style = normal_style
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = parse_length(0)
                apply_line_spacing(p.paragraph_format, body_ls)
                r = p.add_run("Abstract")
                set_run_font(r, east_asia=latin, size_pt=h1_size, bold=True, latin=latin)
            elif re.match(en_abs_re, t):
                past_abstract = True
                if normal_style is not None:
                    p.style = normal_style
                content = re.sub(r"^\s*Abstract\s*[\uff1a:]\s*", "", t, flags=re.I)
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = parse_length(0)
                apply_line_spacing(p.paragraph_format, body_ls)
                r1 = p.add_run("Abstract: ")
                set_run_font(r1, east_asia=latin, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif re.match(en_kw_re, t):
                en_kw_para = p
                normalized = normalize_en_keywords(t) or t
                content = re.sub(r"^\s*Key\s*words\s*[\uff1a:]\s*", "", normalized, flags=re.I)
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = parse_length(0)
                apply_line_spacing(p.paragraph_format, body_ls)
                r1 = p.add_run("Key words: ")
                set_run_font(r1, east_asia=latin, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif has_explicit_en_abstract and not past_abstract and not contains_cjk(t) and not re.match(r"^\s*(Abstract|Key\s*words)\s*[\uff1a:]", t, re.I) and len(t) > 20 and not re.match(r"^[\(\uff08]", t):
                en_title_seen = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = parse_length(0)
                apply_line_spacing(p.paragraph_format, body_ls)
                set_para_runs_font(p, east_asia=latin, size_pt=h1_size, bold=True, latin=latin)
            elif has_explicit_en_abstract and not past_abstract and en_title_seen and not contains_cjk(t) and not re.match(r"^[\(\uff08]", t) and not re.match(r"^\s*(Abstract|Key\s*words)\s*[\uff1a:]", t, re.I):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = parse_length(0)
                apply_line_spacing(p.paragraph_format, body_ls)
                set_para_runs_font(p, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif has_explicit_en_abstract and not past_abstract and re.match(r"^[\(\uff08]", t) and re.search(r"(China|University|College)", t, re.I):
                new_t = t
                if new_t.startswith("("):
                    new_t = "（" + new_t[1:]
                if new_t.endswith(")"):
                    new_t = new_t[:-1] + "）"
                if new_t != t:
                    for run in p.runs:
                        if run.text:
                            run.text = run.text.replace("(", "（").replace(")", "）")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = parse_length(0)
                apply_line_spacing(p.paragraph_format, body_ls)
                set_para_runs_font(p, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            else:
                if contains_cjk(t):
                    if body_align is not None:
                        p.alignment = body_align
                    p.paragraph_format.first_line_indent = body_indent
                    apply_line_spacing(p.paragraph_format, body_ls)
                    set_para_runs_font(p, east_asia=body_font, size_pt=body_size,
                                       bold=False, latin=latin)
                elif t:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = parse_length(0)
                    apply_line_spacing(p.paragraph_format, body_ls)
                    set_para_runs_font(p, east_asia=latin, size_pt=body_size,
                                       bold=False, latin=latin)

        if front_has_english and not has_explicit_en_abstract:
            warnings.append("前置页缺少明确的英文Abstract标题，已保留原文，未强行补写。")

        if cn_kw_para is not None:
            insert_page_break_after(cn_kw_para)
        if en_kw_para is not None:
            insert_page_break_after(en_kw_para)
    def _apply_heading_para(para, align, hcfg, font, size, bold):
        if align is not None:
            para.alignment = align
        para.paragraph_format.first_line_indent = parse_length(0)
        apply_line_spacing(para.paragraph_format, body_ls)
        sb = hcfg.get("space_before", 0)
        sa = hcfg.get("space_after", 0)
        if sb >= 0:
            apply_paragraph_spacing(para.paragraph_format, "before", sb)
        if sa >= 0:
            apply_paragraph_spacing(para.paragraph_format, "after", sa)
        set_para_runs_font(para, east_asia=font, size_pt=size, bold=bold, latin=latin)

    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        level = get_paragraph_heading_level(para)
        t = para.text.strip()
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        sn = para.style.name if para.style else ""

        if level is not None and t_nospace in st_map:
            entry = st_map[t_nospace]
            para.text = entry["display"]
            para.alignment = _ALIGN_MAP.get(entry.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.first_line_indent = parse_length(0)
            apply_line_spacing(para.paragraph_format, body_ls)
            _h1cfg = cfg["headings"]["h1"]
            _sb = _h1cfg.get("space_before", 0)
            _sa = _h1cfg.get("space_after", 0)
            if _sb >= 0:
                apply_paragraph_spacing(para.paragraph_format, "before", _sb)
            if _sa >= 0:
                apply_paragraph_spacing(para.paragraph_format, "after", _sa)
            set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                               bold=True, latin=latin)
        elif level == 1:
            _apply_heading_para(para, h1_align, cfg["headings"]["h1"], h1_font, h1_size, h1_bold)
            if t.startswith("附录"):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)
        elif level == 2:
            _apply_heading_para(para, h2_align, cfg["headings"]["h2"], h2_font, h2_size, h2_bold)
        elif level == 3:
            _apply_heading_para(para, h3_align, cfg["headings"]["h3"], h3_font, h3_size, h3_bold)
        elif level == 4:
            _apply_heading_para(para, h4_align, cfg["headings"]["h4"], h4_font, h4_size, h4_bold)
        if level is not None:
            _ensure_keep_next(para._element)

    ref_cfg = cfg["references"]
    in_refs = False
    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        level = get_paragraph_heading_level(para)
        t = para.text.strip().replace(" ", "").replace("\u3000", "")
        is_h1_like = level == 1
        if is_h1_like and t == ref_key.replace(" ", "").replace("\u3000", ""):
            in_refs = True
            continue
        if is_h1_like and in_refs:
            in_refs = False
        if in_refs and para.text.strip() and not (para.style and is_heading_style(para.style)):
            para.paragraph_format.first_line_indent = parse_length(ref_cfg["first_line_indent"])
            para.paragraph_format.left_indent = parse_length(ref_cfg["left_indent"])
            apply_line_spacing(para.paragraph_format, body_ls)
            set_para_runs_font(para, east_asia=body_font, size_pt=body_size,
                               bold=False, latin=latin)

    cap_cfg, fig_pat, tbl_pat = _format_captions(doc, cfg, preserved_para_ids)

    if cap_cfg.get("check_numbering", True):
        warnings.extend(_check_caption_numbering(doc, fig_pat, tbl_pat, cfg))

    if cap_cfg.get("use_seq_fields", True):
        from . import numbering
        numbering.setup_figure_captions(doc, cfg)
        numbering.setup_table_captions(doc, cfg)

    try:
        warnings.extend(check_citations(doc, cfg))
    except Exception as exc:
        print(f"引用检查出错（已跳过）: {exc}", file=sys.stderr)

    _cite_comma = re.compile(r",\s*((?:19|20)\d{2})")
    for para in doc.paragraphs:
        if _is_preserved_front_para(para):
            continue
        runs = para.runs
        if not runs:
            continue
        full_text = "".join(r.text or "" for r in runs)
        if not _cite_comma.search(full_text):
            continue
        pos = 0
        run_spans = []
        for r in runs:
            rt = r.text or ""
            run_spans.append((r, pos, pos + len(rt)))
            pos += len(rt)
        for m in _cite_comma.finditer(full_text):
            insert_pos = m.start() + 1
            for r, start, end in run_spans:
                if start <= insert_pos < end:
                    local = insert_pos - start
                    rt = r.text or ""
                    r.text = rt[:local] + " " + rt[local:]
                    break
            print(f"  引用逗号修正: \"{m.group(0)}\" → \", {m.group(1)}\"")

    try:
        apply_ref_crosslinks(doc, cfg)
    except Exception as exc:
        print(f"交叉引用创建出错（已跳过）: {exc}", file=sys.stderr)

    _format_tables(doc, cfg)
    if toc_enabled:
        abstract_demotions = demote_abstract_heading_styles(doc, cfg, include_abstract=toc_cfg.get("exclude_abstract_headings", True))
        if abstract_demotions:
            print(f"以下段落已从 Heading 样式解除，不参与目录 ({len(abstract_demotions)} 个):", file=sys.stderr)
            for item in abstract_demotions:
                print(item, file=sys.stderr)
    if toc_enabled:
        insert_toc(doc, cfg)
    toc_match = _find_special_display(cfg, "目录", raw=True)
    first_body_h1 = find_first_body_heading(doc, cfg)
    body_started = first_body_h1 is None
    for para in doc.paragraphs:
        if get_paragraph_heading_level(para) != 1:
            continue
        if para is first_body_h1:
            body_started = True
            continue
        if not body_started:
            continue
        t = para.text.strip().replace(" ", "").replace("\u3000", "")
        if t == toc_match:
            continue
        para.paragraph_format.page_break_before = True

    custom_cover = cfg.get("cover", {}).get("custom_docx", "")
    use_custom_cover = bool(custom_cover and os.path.isfile(custom_cover))
    if use_custom_cover:
        insert_cover_and_declaration(doc, cfg, config_path, skip_cover=True)
    elif cfg["cover"]["enabled"] and not _has_cover(doc, cfg):
        insert_cover_and_declaration(doc, cfg, config_path)

    _restore_preserved_front_paragraphs()
    if toc_enabled:
        demote_abstract_heading_styles(doc, cfg, include_abstract=toc_cfg.get("exclude_abstract_headings", True), aggressive_body_demote=False)
    setup_page_numbers(doc, cfg)
    try:
        setup_headers(doc, cfg)
    except Exception as e:
        print(f"  [警告] 页眉设置出错，已跳过: {e}", file=sys.stderr)
    doc.save(output_path)
    if use_custom_cover:
        success, err = _insert_cover_via_vbs(output_path, custom_cover)
        if success:
            cfg.setdefault("_runtime", {})["custom_cover_sections"] = 1
            print("自定义封面已插入 (VBS)", file=sys.stderr)
        else:
            print(f"自定义封面插入失败（VBS不可用，已跳过）: {err}", file=sys.stderr)
    patch_theme_fonts(output_path, cfg)
    if renum_changes:
        warnings.append("标题编号已自动修正:")
        warnings.extend(renum_changes)
    return warnings


def patch_theme_fonts(docx_path, cfg):
    import xml.etree.ElementTree as ET
    theme = cfg.get("theme_fonts", {})
    theme_latin = theme.get("latin", "Times New Roman")
    theme_hans = theme.get("hans", "宋体")

    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ET.register_namespace("a", a_ns)
    ns = {"a": a_ns}

    fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=os.path.dirname(docx_path))
    os.close(fd)
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/theme/theme1.xml":
                    root = ET.fromstring(data)
                    for minor in root.findall(".//a:minorFont", ns):
                        lat = minor.find("a:latin", ns)
                        if lat is not None:
                            lat.set("typeface", theme_latin)
                        for font in minor.findall("a:font", ns):
                            if font.get("script") == "Hans":
                                font.set("typeface", theme_hans)
                    for major in root.findall(".//a:majorFont", ns):
                        lat = major.find("a:latin", ns)
                        if lat is not None:
                            lat.set("typeface", theme_latin)
                        for font in major.findall("a:font", ns):
                            if font.get("script") == "Hans":
                                font.set("typeface", theme_hans)
                    data = ET.tostring(root, encoding="unicode").encode("utf-8")
                zout.writestr(item, data)
        os.replace(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Universal thesis formatter")
    parser.add_argument("--input", required=True, help="Input docx")
    parser.add_argument("--output", required=True, help="Output docx")
    parser.add_argument("--config", help="Path to thesis_config.yaml")
    args = parser.parse_args()

    cfg, cfg_path = resolve_config(cli_config=args.config, input_path=args.input)
    apply_format(args.input, args.output, config=cfg, config_path=cfg_path)
    print(f"OK {args.output}")







