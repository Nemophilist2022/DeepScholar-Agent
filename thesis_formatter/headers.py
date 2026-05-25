from ._common import _ALIGN_MAP, set_run_font, _HEADING_STYLE_IDS, parse_length
from .page import _set_even_odd_on_doc, get_body_start_section_index
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_header_border(paragraph, width_pt=0.75, style="single"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    sz = int(width_pt * 8)
    val = "double" if style == "double" else "single"
    bottom.set(qn("w:val"), val)
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)


def _render_header_para(paragraph, text, cfg, hf_cfg, doc, align="center"):
    paragraph.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.space_before = parse_length(0)
    paragraph.paragraph_format.space_after = parse_length(0)
    rendered = text
    if "{chapter_title}" in rendered:
        rendered_prefix, rendered_suffix = rendered.split("{chapter_title}", 1)
        paragraph.clear()

        h1_style_id = _HEADING_STYLE_IDS.get(1, "Heading1")
        for sty in doc.styles:
            if sty.style_id == h1_style_id:
                name_el = sty.element.find(qn("w:name"))
                if name_el is not None:
                    aliases_el = sty.element.find(qn("w:aliases"))
                    if aliases_el is None:
                        aliases_el = OxmlElement("w:aliases")
                        name_el.addnext(aliases_el)
                    existing = aliases_el.get(qn("w:val"), "")
                    if "标题 1" not in existing:
                        new_val = ("标题 1," + existing) if existing else "标题 1"
                        aliases_el.set(qn("w:val"), new_val)
                break
        h1_style_name = "标题 1"

        def _hf_run(txt=None):
            r = paragraph.add_run(txt) if txt else paragraph.add_run()
            set_run_font(r, east_asia=hf_cfg["font"],
                         size_pt=parse_length(hf_cfg["font_size"]),
                         bold=hf_cfg.get("bold", False),
                         latin=cfg["fonts"]["latin"])
            return r

        if rendered_prefix:
            _hf_run(rendered_prefix)
        r1 = _hf_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        r1._element.append(begin)
        r2 = _hf_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' STYLEREF "{h1_style_name}" '
        r2._element.append(instr)
        r3 = _hf_run()
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        r3._element.append(sep)
        _hf_run("(章标题)")
        r5 = _hf_run()
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r5._element.append(end)
        if rendered_suffix:
            _hf_run(rendered_suffix)
    else:
        paragraph.clear()
        r = paragraph.add_run(rendered)
        set_run_font(r, east_asia=hf_cfg["font"],
                     size_pt=parse_length(hf_cfg["font_size"]),
                     bold=hf_cfg.get("bold", False),
                     latin=cfg["fonts"]["latin"])
    if hf_cfg.get("border_bottom", False):
        _add_header_border(paragraph,
                           hf_cfg.get("border_bottom_width", 0.75),
                           hf_cfg.get("border_bottom_style", "single"))


def _clear_header_story(story):
    story.is_linked_to_previous = False
    for p in story.paragraphs:
        p.clear()


def _clear_section_headers(section):
    for story in (section.header, section.first_page_header, section.even_page_header):
        _clear_header_story(story)

def setup_headers(doc, cfg):
    hf = cfg.get("header_footer", {})
    if not hf.get("enabled", False):
        return
    scope = hf.get("scope", "body")
    diff_oe = hf.get("different_odd_even", True)
    first_no = hf.get("first_page_no_header", False)
    odd_align = hf.get("odd_page_align", "center")
    even_align = hf.get("even_page_align", "center")

    if diff_oe:
        _set_even_odd_on_doc(doc)
    doc_has_even_odd = doc.settings.element.find(qn("w:evenAndOddHeaders")) is not None
    cover_sections = max(0, int(cfg.get("_runtime", {}).get("custom_cover_sections", 0) or 0))
    body_section_index = get_body_start_section_index(doc, cfg) if len(doc.sections) > 1 else 0

    for idx, section in enumerate(doc.sections):
        if idx < cover_sections:
            _clear_section_headers(section)
            continue
        is_front = idx < body_section_index
        if scope == "body" and is_front:
            _clear_section_headers(section)
            continue

        if first_no:
            sect_pr = section._sectPr
            if sect_pr.find(qn("w:titlePg")) is None:
                sect_pr.append(OxmlElement("w:titlePg"))

        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        _render_header_para(hp, hf["odd_page_text"], cfg, hf, doc, align=odd_align)

        if doc_has_even_odd:
            even_header = section.even_page_header
            even_header.is_linked_to_previous = False
            for p in even_header.paragraphs:
                p.clear()
            ep = even_header.paragraphs[0] if even_header.paragraphs else even_header.add_paragraph()
            even_text = hf["even_page_text"] if diff_oe else hf["odd_page_text"]
            even_a = even_align if diff_oe else odd_align
            _render_header_para(ep, even_text, cfg, hf, doc, align=even_a)

        if first_no:
            first_header = section.first_page_header
            first_header.is_linked_to_previous = False
            for p in first_header.paragraphs:
                p.clear()
            first_footer = section.first_page_footer
            first_footer.is_linked_to_previous = False

