import copy
import os
import re
import sys

from thesis_config import resolve_logo_path
from ._common import parse_length, get_paragraph_heading_level
from docx import Document
from docx.opc.part import Part as OpcPart
from docx.opc.packuri import PackURI
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def insert_custom_cover_via_vbs(target_doc_path, cover_path, output_path=None, vbs_path=None):
    """
    使用 VBS 调用 Word COM 接口插入封面，保留完整格式。
    不需要安装 pywin32，只需要 Microsoft Word。

    Args:
        target_doc_path: 目标文档路径
        cover_path: 封面文档路径
        output_path: 输出文档路径（默认覆盖原文档）
        vbs_path: VBS 脚本路径（默认使用同目录下的 InsertCover.vbs）

    Returns:
        bool: 成功返回 True，失败返回 False
    """
    import subprocess

    if output_path is None:
        output_path = target_doc_path

    if vbs_path is None:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        vbs_path = os.path.join(script_dir, "..", "InsertCover.vbs")
        vbs_path = os.path.abspath(vbs_path)

    if not os.path.exists(vbs_path):
        print(f"警告: VBS 脚本不存在: {vbs_path}")
        return False

    try:
        result = subprocess.run(
            ["cscript", "//NoLogo", vbs_path, target_doc_path, cover_path, output_path],
            capture_output=True,
            text=True,
            encoding="gbk",
            timeout=60
        )
        if result.returncode == 0:
            return True
        else:
            print(f"VBS 执行失败: {result.stderr}")
            return False
    except FileNotFoundError:
        print("警告: Windows cscript 不可用")
        return False
    except subprocess.TimeoutExpired:
        print("警告: Word 操作超时")
        return False
    except Exception as e:
        print(f"VBS 调用失败: {e}")
        return False


def insert_custom_cover_via_compose(target_doc_path, cover_path, output_path):
    """
    使用 docxcompose 合并文档，保留更多格式信息。

    需要安装: pip install docxcompose

    Args:
        target_doc_path: 目标文档路径
        cover_path: 封面文档路径
        output_path: 输出文档路径

    Returns:
        bool: 成功返回 True，失败返回 False
    """
    try:
        from docxcompose.composer import Composer

        target = Document(target_doc_path)
        cover = Document(cover_path)

        composer = Composer(target)
        composer.append(cover)

        composer.save(output_path)
        return True
    except ImportError:
        print("警告: 未安装 docxcompose，可使用 'pip install docxcompose' 安装")
        return False
    except Exception as e:
        print(f"docxcompose 合并失败: {e}")
        return False


def _has_cover(doc, cfg, scan_limit=30):
    cover_title = cfg["cover"].get("title_text", "毕业论文")
    keywords = ["毕业论文", "毕业设计"]
    if cover_title:
        keywords.append(cover_title.replace(" ", ""))
    for para in doc.paragraphs[:scan_limit]:
        t = para.text.replace(" ", "").replace("\u3000", "")
        if any(kw in t for kw in keywords):
            return True
    return False




def find_existing_cover_end(doc, cfg, scan_limit=80):
    """Return the paragraph index where an existing cover page likely ends."""
    if not _has_cover(doc, cfg, scan_limit=scan_limit):
        return 0

    sec = cfg.get("sections", {})
    cn_kw_re = sec.get("cn_keywords_pattern", r"^\s*关键词\s*[：:]")
    en_abs_re = sec.get("en_abstract_pattern", r"(?i)^\s*Abstract\s*[：:]")

    special_titles = {
        st.get("match", "").replace(" ", "").replace("\u3000", "")
        for st in cfg.get("special_titles", [])
        if st.get("match")
    }
    declaration_titles = {
        decl.get("title", "").replace(" ", "").replace("\u3000", "")
        for decl in cfg.get("declarations", [])
        if decl.get("title")
    }

    cover_title = cfg.get("cover", {}).get("title_text", "毕业论文")
    cover_keywords = ["毕业论文", "毕业设计"]
    if cover_title:
        cover_keywords.append(cover_title.replace(" ", ""))

    saw_cover = False
    limit = min(len(doc.paragraphs), scan_limit)
    for idx, para in enumerate(doc.paragraphs[:limit]):
        text = para.text.strip()
        normalized = text.replace(" ", "").replace("\u3000", "")

        if normalized and any(keyword and keyword in normalized for keyword in cover_keywords):
            saw_cover = True

        if not saw_cover or idx == 0 or not text:
            continue

        if get_paragraph_heading_level(para) == 1:
            return idx
        if normalized in special_titles or normalized in declaration_titles or normalized == "摘要":
            return idx
        if re.match(cn_kw_re, text) or re.match(en_abs_re, text):
            return idx

    return limit

def insert_custom_cover(doc, cover_path, use_word_com=True):
    if use_word_com:
        import tempfile

        temp_target = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        temp_output = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        temp_target.close()
        temp_output.close()
        temp_target_name = temp_target.name
        temp_output_name = temp_output.name

        try:
            doc.save(temp_target_name)

            success = insert_custom_cover_via_vbs(
                temp_target_name, cover_path, temp_output_name
            )

            if success:
                from docx import Document as Doc
                merged_doc = Doc(temp_output_name)

                doc.element.body.clear()
                for el in merged_doc.element.body:
                    doc.element.body.append(copy.deepcopy(el))
                return
        except Exception as e:
            print(f"Word 方法失败，回退到 XML 手动复制: {e}")
        finally:
            for path in (temp_target_name, temp_output_name):
                try:
                    os.unlink(path)
                except OSError:
                    pass

    # 回退方案：使用 python-docx XML 手动复制
    cover_doc = Document(cover_path)

    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    main_styles = doc.styles.element
    existing_ids = {
        s.get(qn("w:styleId"))
        for s in main_styles.findall(f"{{{ns_w}}}style")
    }
    for style_el in cover_doc.styles.element.findall(f"{{{ns_w}}}style"):
        sid = style_el.get(qn("w:styleId"))
        if sid and sid not in existing_ids:
            main_styles.append(copy.deepcopy(style_el))
            existing_ids.add(sid)

    try:
        cover_num_part = cover_doc.part.numbering_part
        main_num_part = doc.part.numbering_part
        if cover_num_part is not None and main_num_part is not None:
            ns = {"w": ns_w}
            for anum in cover_num_part.element.findall("w:abstractNum", ns):
                main_num_part.element.append(copy.deepcopy(anum))
            for num in cover_num_part.element.findall("w:num", ns):
                main_num_part.element.append(copy.deepcopy(num))
    except Exception as e:
        print(f"[警告] 编号部件合并失败，已跳过: {e}", file=sys.stderr)

    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    max_img = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            m = re.search(r"image(\d+)", str(rel.target_part.partname))
            if m:
                max_img = max(max_img, int(m.group(1)))
    next_img = max_img + 1

    rid_map = {}
    for rel in cover_doc.part.rels.values():
        try:
            if "image" in rel.reltype:
                src = rel.target_part
                ext = os.path.splitext(str(src.partname))[1] or ".png"
                new_name = PackURI(f"/word/media/image{next_img}{ext}")
                next_img += 1
                img_part = OpcPart(
                    new_name, src.content_type, bytes(src.blob), doc.part.package
                )
                new_rId = doc.part.relate_to(img_part, rel.reltype)
                rid_map[rel.rId] = new_rId
            elif rel.is_external:
                new_rId = doc.part.rels.get_or_add_ext_rel(
                    rel.reltype, rel.target_ref
                )
                rid_map[rel.rId] = new_rId
        except Exception as e:
            print(f"[警告] 封面资源引用合并失败，已跳过: {e}", file=sys.stderr)

    _strip_tags = {
        qn("w:commentRangeStart"), qn("w:commentRangeEnd"),
        qn("w:commentReference"), qn("w:annotationRef"),
        qn("w:footnoteReference"), qn("w:endnoteReference"),
    }

    main_body = doc.element.body
    first_child = main_body[0] if len(main_body) > 0 else None
    insert_idx = list(main_body).index(first_child) if first_child is not None else len(main_body)

    for el in list(cover_doc.element.body):
        if el.tag == qn("w:sectPr"):
            continue
        el_copy = copy.deepcopy(el)

        for bad in list(el_copy.iter()):
            if bad.tag in _strip_tags:
                parent = bad.getparent()
                if parent is not None:
                    parent.remove(bad)

        for node in el_copy.iter():
            for attr_name, attr_val in list(node.attrib.items()):
                if attr_val in rid_map:
                    node.set(attr_name, rid_map[attr_val])

        main_body.insert(insert_idx, el_copy)
        insert_idx += 1

    brk_p = OxmlElement("w:p")
    brk_r = OxmlElement("w:r")
    brk_br = OxmlElement("w:br")
    brk_br.set(qn("w:type"), "page")
    brk_r.append(brk_br)
    brk_p.append(brk_r)
    main_body.insert(insert_idx, brk_p)


def insert_cover_and_declaration(doc, cfg, config_path=None, skip_cover=False):
    cover = cfg["cover"]
    latin = cfg["fonts"]["latin"]

    def mk_run(text, ea="宋体", sz_hp=None, bold=False, uline=False):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), latin)
        rf.set(qn("w:hAnsi"), latin)
        rf.set(qn("w:eastAsia"), ea)
        rPr.append(rf)
        if sz_hp:
            for tag in ("w:sz", "w:szCs"):
                s = OxmlElement(tag)
                s.set(qn("w:val"), str(sz_hp))
                rPr.append(s)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if uline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        return r

    def mk_para(runs=None, align=None, fi=None, fi_chars=None,
                ls_auto=None, ls_exact=None, sb=None, sa=None):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        if align:
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), align)
            pPr.append(jc)
        sp = OxmlElement("w:spacing")
        need_sp = False
        if ls_exact is not None:
            sp.set(qn("w:line"), str(ls_exact))
            sp.set(qn("w:lineRule"), "exact")
            need_sp = True
        elif ls_auto is not None:
            sp.set(qn("w:line"), str(ls_auto))
            sp.set(qn("w:lineRule"), "auto")
            need_sp = True
        if sb is not None:
            sp.set(qn("w:before"), str(sb))
            need_sp = True
        if sa is not None:
            sp.set(qn("w:after"), str(sa))
            need_sp = True
        if need_sp:
            pPr.append(sp)
        if fi is not None or fi_chars is not None:
            ind = OxmlElement("w:ind")
            if fi is not None:
                ind.set(qn("w:firstLine"), str(fi))
            if fi_chars is not None:
                ind.set(qn("w:firstLineChars"), str(fi_chars))
            pPr.append(ind)
        p.append(pPr)
        for r in (runs or []):
            p.append(r)
        return p

    NBSP = "\u00a0"

    def mk_field(label, uline_chars=33):
        return mk_para([
            mk_run(label, sz_hp=30, bold=True),
            mk_run(" ", sz_hp=30),
            mk_run(NBSP * uline_chars, sz_hp=30, uline=True),
        ], fi_chars=400, ls_exact=700, sa=0, sb=0)

    elements = []

    if not skip_cover:
        title_sz_hp = int(cover["title_font_size"] * 2)
        thesis_sz_hp = int(cover["thesis_title_size"] * 2)
        thesis_font = cover["thesis_title_font"]

        elements.append(mk_para(align="center", ls_auto=360))
        elements.append(mk_para(
            [mk_run(cover["title_text"], sz_hp=title_sz_hp, bold=True)],
            align="center", sa=161,
        ))
        elements.append(mk_para(align="center", ls_auto=360))
        elements.append(mk_para(
            [mk_run(cover["thesis_title_placeholder"], ea=thesis_font,
                    sz_hp=thesis_sz_hp, bold=True)],
            align="center", sa=161,
        ))
        for sz in (30, 30, 21, 21, 30, 30):
            elements.append(mk_para([mk_run(" ", sz_hp=sz)], ls_auto=360, sa=0, sb=0))

        for field in cover["fields"]:
            elements.append(mk_field(field["label"], field["underline_chars"]))

        adv = cover["advisor"]
        elements.append(mk_para([
            mk_run(adv["label"], sz_hp=30, bold=True),
            mk_run(" ", sz_hp=30),
            mk_run(NBSP * adv["underline_chars"], sz_hp=30, uline=True),
            mk_run(" ", sz_hp=30),
            mk_run(adv["title_label"], sz_hp=30, bold=True),
            mk_run(" ", sz_hp=30),
            mk_run(NBSP * adv["title_underline_chars"], sz_hp=30, uline=True),
        ], fi_chars=400, ls_exact=700, sa=0, sb=0))

        dt = cover["date"]
        date_runs = [mk_run(dt["label"], sz_hp=30, bold=True)]
        for seg in dt["segments"]:
            date_runs.extend([
                mk_run(" ", sz_hp=30),
                mk_run(NBSP * dt["segment_underline_chars"], sz_hp=30, uline=True),
                mk_run(" ", sz_hp=30),
                mk_run(seg, sz_hp=30, bold=True),
            ])
        elements.append(mk_para(date_runs, fi_chars=400, ls_exact=700, sa=0, sb=0))

        elements.append(mk_para(align="center", sa=161))

        pb = OxmlElement("w:p")
        pb_r = OxmlElement("w:r")
        pb_br = OxmlElement("w:br")
        pb_br.set(qn("w:type"), "page")
        pb_r.append(pb_br)
        pb.append(pb_r)
        elements.append(pb)

    declarations = cfg.get("declarations", [])
    if declarations:
        elements.append(mk_para(ls_auto=360))

    for idx, decl in enumerate(declarations):
        h1_font = cfg["fonts"]["h1"]
        elements.append(mk_para(
            [mk_run(decl["title"], ea=h1_font, sz_hp=28)],
            align="center", sb=161, sa=161, ls_auto=360,
        ))
        elements.append(mk_para([mk_run(
            decl["body"], sz_hp=24,
        )], fi=540, ls_auto=360))
        elements.append(mk_para(fi=540, ls_auto=360))
        elements.append(mk_para([mk_run(
            decl["signature"], sz_hp=24,
        )], fi=540, ls_auto=360))

        if "date_line" in decl:
            last_el = mk_para([mk_run(
                decl["date_line"], sz_hp=24,
            )], fi=540, ls_auto=360)
            elements.append(last_el)
        else:
            last_el = elements[-1]

        if idx < len(declarations) - 1:
            for _ in range(4):
                elements.append(mk_para(ls_auto=360))
        elif idx == len(declarations) - 1:
            elements.append(mk_para(fi=540, ls_auto=360))

    if elements:
        final_el = elements[-1]
        pPr = final_el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            final_el.insert(0, pPr)
        sect_pr = OxmlElement("w:sectPr")
        existing_sect = doc.sections[0]._sectPr
        for attr_name in ("pgSz", "pgMar"):
            src = existing_sect.find(qn(f"w:{attr_name}"))
            if src is not None:
                sect_pr.append(copy.deepcopy(src))
        pPr.append(sect_pr)

    body = doc.element.body
    for el in reversed(elements):
        body.insert(0, el)

    if len(doc.sections) > 0:
        s0_footer = doc.sections[0].footer
        s0_footer.is_linked_to_previous = False
        for p in s0_footer.paragraphs:
            p.clear()
    if len(doc.sections) > 1:
        doc.sections[1].footer.is_linked_to_previous = False

    if not skip_cover:
        logo_path = resolve_logo_path(cfg, config_path)
        if logo_path:
            p0_para = doc.paragraphs[0]
            run = p0_para.add_run()
            run.add_picture(logo_path,
                            width=parse_length(cover["logo_width_pt"]),
                            height=parse_length(cover["logo_height_pt"]))

    if len(doc.sections) > 1:
        doc.sections[1].footer.is_linked_to_previous = False
